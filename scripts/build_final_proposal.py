from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\Project\lariska-ai\docs\LARISKA_AI_Proposal_Final_Competition.docx")

NAVY = "172554"
PURPLE = "715BC9"
MUTED = "475569"
LIGHT = "F5F3FF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_table_geometry(table, widths=None):
    """Tetapkan geometri DXA agar tabel stabil di Word/LibreOffice."""
    widths = widths or [6.5]
    width_dxa = [int(round(w * 1440)) for w in widths]
    total_dxa = sum(width_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col in list(grid):
        grid.remove(col)
    for width in width_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        set_cell_shading(cell, PURPLE)
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cells[i])
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1E293B")
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in (("Heading 1", 16, NAVY), ("Heading 2", 13, PURPLE), ("Heading 3", 11.5, NAVY)):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "LARISKA AI | Proposal Inovasi"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(section.footer.paragraphs[0])

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LARISKA AI")
    r.font.name = "Calibri"; r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sales Brain untuk UMKM Indonesia")
    r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(PURPLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Platform penjualan WhatsApp yang mengubah percakapan pelanggan menjadi keputusan bisnis yang aman, kontekstual, dan dapat diaudit.").italic = True
    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(callout.rows[0])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT); set_cell_margins(cell, 180, 120, 180, 120)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cp.add_run("Posisi inti: LARISKA bukan chatbot penjawab pesan. LARISKA adalah Sales Brain yang memisahkan keputusan bisnis deterministik dari bahasa natural.")
    rr.font.bold = True; rr.font.color.rgb = RGBColor.from_string(NAVY)
    set_table_geometry(callout, [6.5])
    doc.add_paragraph()
    for line in ("Diajukan oleh: Kelompok iki", "AI Innovation Challenge - COMPFEST 18", "2026"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(line)
    doc.add_page_break()

    def h(text, level=1):
        doc.add_heading(text, level=level)
    def para(text, bold_prefix=None):
        p = doc.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            p.add_run(bold_prefix).bold = True; p.add_run(text[len(bold_prefix):])
        else:
            p.add_run(text)
        return p
    def bullets(items):
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    h("Ringkasan Eksekutif")
    para("LARISKA AI adalah Sales Brain untuk UMKM Indonesia yang melayani percakapan pelanggan melalui WhatsApp. Sistem ini menerima teks maupun voice note, memahami maksud pembeli, mengingat konteks produk dan percakapan, lalu menjalankan keputusan negosiasi secara deterministik sebelum menyusun balasan Bahasa Indonesia yang natural.")
    para("Inovasi utama LARISKA berada pada pemisahan tegas antara lapisan keputusan dan lapisan bahasa. LightGBM memberi rekomendasi aksi negosiasi dari sinyal margin, stok, loyalitas, diskon yang diminta, dan waktu; hard guardrail Python memastikan harga tidak pernah melewati floor price. Gemini hanya digunakan untuk ekstraksi terstruktur dan perumusan bahasa, sehingga tidak dapat mengubah keputusan harga.")
    para("MVP telah terintegrasi dengan WhatsApp Cloud API, Supabase, Midtrans Sandbox, Whisper untuk voice note, dashboard operasional, pencatatan order dan pembayaran, serta pengurangan stok yang terlacak. Dengan demikian, LARISKA menyelesaikan funnel penjualan dari discovery produk hingga status pemenuhan pesanan.")
    para("Proposal ini membedakan secara tegas antara kemampuan yang sudah diimplementasikan dan roadmap. Dengan begitu, LARISKA dapat dinilai sebagai MVP yang dapat diuji saat ini sekaligus memiliki arah pengembangan yang bertanggung jawab untuk pemakaian UMKM di lapangan.")

    h("1. Latar Belakang dan Rumusan Masalah")
    para("Industri mikro dan kecil memiliki peran penting dalam penciptaan kerja dan penyerapan tenaga kerja lokal, namun masih menghadapi keterbatasan manajemen serta kualitas SDM (BPS, 2025). Pada level operasional, banyak pemilik usaha masih melayani katalog, stok, tawar-menawar, pembayaran, dan pertanyaan purnajual melalui WhatsApp secara manual sambil mengurus produksi maupun pengiriman. Kondisi ini membuat respons terlambat, konteks pelanggan mudah hilang, dan negosiasi tidak konsisten.")
    bullets([
        "Pelanggan mengirim chat bersamaan, sering memakai typo, bahasa informal, dan voice note.",
        "Pemilik usaha sulit menjaga harga tetap kompetitif tanpa mengorbankan margin atau melanggar batas harga aman.",
        "Informasi stok, spesifikasi, dan status pesanan tersebar sehingga jawaban manual mudah tidak konsisten.",
        "Chatbot template hanya memetakan kata kunci; ia tidak menyatukan konteks percakapan, keputusan harga, pembayaran, dan operasi bisnis."
    ])
    h("1.1 Mengapa Sales Brain diperlukan", 2)
    para("Masalah utama bukan hanya volume chat. Pada satu percakapan, pemilik UMKM perlu menyampaikan fakta katalog yang konsisten, menangkap maksud pembeli, menilai tawaran tanpa mengorbankan margin, memastikan stok, mengarahkan checkout, dan tetap mampu menindaklanjuti pesanan setelah pembayaran. Kesalahan pada salah satu titik dapat menurunkan kepercayaan pelanggan atau merugikan usaha. Karena itu LARISKA dirancang sebagai lapisan keputusan penjualan yang terhubung ke operasi, bukan auto-reply umum.")
    para("Hubungan dengan tema backbone ekonomi bersifat langsung: ketika usaha mikro dan kecil dapat merespons permintaan dengan cepat, menjaga margin, dan memproses pesanan secara tertib, kapasitas layanan pemilik usaha tidak lagi semata dibatasi oleh waktu membalas chat. LARISKA berfokus pada penguatan proses penjualan unit usaha, bukan membuat klaim makroekonomi yang tidak dapat diuji oleh MVP.")
    h("1.2 Rumusan masalah", 2)
    for item in [
        "Bagaimana mengubah percakapan WhatsApp yang informal, ber-typo, atau berbentuk voice note menjadi data transaksi yang dapat ditindaklanjuti?",
        "Bagaimana memberikan negosiasi yang fleksibel dan menyenangkan tanpa memberi diskon di bawah batas aman UMKM?",
        "Bagaimana membuat informasi produk, stok, pembayaran, dan pengiriman dapat ditelusuri dalam satu alur operasional?"
    ]:
        doc.add_paragraph(item, style="List Number")

    h("2. Tujuan, Pengguna, dan Nilai Manfaat")
    para("Tujuan LARISKA AI adalah memberi UMKM tenaga penjual digital yang dapat bekerja sepanjang waktu tanpa menghilangkan kontrol pemilik usaha atas harga, stok, dan pengalaman pelanggan.")
    add_table(doc, ["Pemangku kepentingan", "Nilai yang diberikan"], [
        ["Pemilik UMKM", "Respons yang konsisten, pengamanan floor price, dashboard produk/order/pembayaran, dan jejak audit keputusan."],
        ["Pelanggan", "Jawaban cepat untuk katalog dan stok, navigasi kategori WhatsApp, negosiasi yang hangat, serta checkout yang jelas."],
        ["Tim operasional", "Order paid dapat dilanjutkan ke shipped dan completed tanpa mengubah status pembayaran secara manual."],
        ["Ekosistem UMKM", "Data interaksi yang terstruktur untuk evaluasi dan retraining batch yang berizin pada masa depan."]
    ], [1.55, 4.95])
    h("2.1 Sasaran keberhasilan MVP", 2)
    bullets([
        "Respons katalog dan stok berasal dari database aktif, bukan pengetahuan bebas model bahasa.",
        "Setiap keputusan negosiasi mempunyai harga/floor-price snapshot yang dapat diaudit.",
        "Pembayaran yang tervalidasi webhook mengubah order menjadi paid dan memicu stok secara konsisten.",
        "Operator dapat membedakan pesan yang dapat ditangani otomatis dari percakapan yang perlu handover."
    ])

    h("3. Solusi dan Alur Pengalaman Pengguna")
    para("Pengalaman pelanggan dirancang sebagai funnel yang singkat dan dapat didemonstrasikan tanpa mengandalkan jawaban bebas LLM untuk fakta bisnis.")
    for i, step in enumerate([
        "Pelanggan mengirim salam; LARISKA menyapa dan mengirim daftar kategori yang dibangun dari katalog aktif Supabase.",
        "Pelanggan memilih kategori atau menyebut produk; LARISKA mengambil rekomendasi, detail, harga, dan stok dari sumber katalog yang sama.",
        "Pelanggan bertanya lewat teks atau voice note; Whisper mentranskripsikan suara Bahasa Indonesia, lalu NLU mengekstrak intent dan entitas.",
        "Saat negosiasi, Sales Brain menghitung keputusan harga dari fitur bisnis dan guardrail mengunci batas aman.",
        "Saat checkout, sistem membuat order, reservasi stok, tautan pembayaran QRIS Midtrans, lalu webhook tervalidasi mengubah pembayaran menjadi success dan order menjadi paid.",
        "Dashboard membantu operator menindaklanjuti paid -> shipped -> completed serta memantau katalog, pelanggan, negosiasi, dan transaksi."
    ], start=1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    h("3.1 Skenario penggunaan prioritas", 2)
    add_table(doc, ["Skenario", "Respons sistem", "Kontrol bisnis"], [
        ["Discovery", "Salam membuka daftar kategori; produk, spesifikasi, harga, dan stok dibaca dari katalog.", "Produk nonaktif/tidak tersedia tidak direkomendasikan."],
        ["Voice note", "Audio WhatsApp ditranskripsi sebelum diproses sebagai pesan pelanggan.", "Hasil tetap melewati deteksi intent dan konteks produk."],
        ["Negosiasi", "Model memberi rekomendasi aksi, lalu sistem menyusun jawaban yang hangat dan persuasif.", "Floor price dan ketersediaan stok tidak dapat dilanggar."],
        ["Checkout", "Order dan payment link dibuat dari harga server-side.", "Status paid hanya berasal dari webhook pembayaran tervalidasi."],
        ["Purnajual", "Operator memperbarui fulfilled flow dari paid ke shipped lalu completed.", "Tidak ada tombol dashboard untuk memalsukan status paid."]
    ], [1.25, 3.3, 1.95])

    h("4. Arsitektur Sistem dan Peran Model")
    para("Arsitektur LARISKA menggunakan pipeline hibrida. Setiap komponen memiliki peran terbatas agar sistem dapat dijelaskan, diuji, dan dipulihkan saat layanan eksternal tidak tersedia.")
    add_table(doc, ["Tahap", "Komponen", "Tanggung jawab"], [
        ["1", "WhatsApp Cloud API", "Menerima teks, interactive reply, voice note, serta delivery status (Meta, 2026)."],
        ["2", "Whisper small + fallback", "Transkripsi voice note Bahasa Indonesia; fallback base/Gemini bila model utama tidak tersedia (Radford et al., 2023)."],
        ["3", "Gemini structured output", "Ekstraksi intent dan entitas: produk, jumlah, penawaran, kategori."],
        ["4", "State Tracking + Supabase", "Menyimpan customer, conversation, message, produk yang dibahas, dan kuantitas terakhir."],
        ["5", "LightGBM + Python guardrail", "Merekomendasikan aksi hold_price/discount/bonus/counter_offer; mengunci floor price dan stok."],
        ["6", "Response Generator", "Menghasilkan bahasa natural; fakta katalog, rekomendasi, dan harga menggunakan data/keputusan upstream."],
        ["7", "Midtrans + Dashboard", "Mencatat payment event, reservasi/pengurangan stok, serta status pemenuhan pesanan."]
    ], [0.45, 1.75, 4.3])
    para("Prinsip keselamatan: Gemini tidak membuat keputusan harga. Untuk pertanyaan stok, harga, spesifikasi, dan rekomendasi kategori, LARISKA merespons secara deterministik dari katalog agar tidak mengarang informasi.")
    h("4.1 Kontrak keputusan", 2)
    para("Pemisahan peran adalah inti arsitektur. NLU menghasilkan struktur intent dan entitas, Emotion Classifier memberi sinyal gaya/urgensi, LightGBM menghasilkan rekomendasi aksi, dan guardrail mengesahkan harga. Response Generator hanya menerima payload keputusan yang telah disetujui. Dengan kontrak ini, kegagalan atau variasi keluaran LLM tidak dapat secara langsung mengubah floor price maupun stok.")
    add_table(doc, ["Lapisan", "Boleh dilakukan", "Tidak boleh dilakukan"], [
        ["LLM / Gemini", "Ekstraksi entitas dan penyusunan bahasa Indonesia yang empatik.", "Menetapkan harga akhir, melampaui floor price, atau mengubah stok."],
        ["LightGBM", "Memberikan rekomendasi aksi dari enam fitur tabular.", "Menjadi satu-satunya otoritas harga."],
        ["Python guardrail", "Mengunci batas harga, stok, dan tindakan fallback.", "Mengabaikan batas produk yang disetel pemilik."],
        ["Supabase", "Menjadi sumber katalog, state pelanggan, order, payment, dan audit log.", "Menyimpan secret API pada data publik."]
    ], [1.35, 2.65, 2.5])
    h("4.2 Alasan pemilihan model", 2)
    add_table(doc, ["Alternatif", "Peran pada LARISKA", "Alasan keputusan desain"], [
        ["Rule deterministik", "Guardrail final untuk harga dan stok.", "Paling dapat diaudit untuk aturan yang tidak boleh dilanggar; bukan pengganti seluruh adaptasi negosiasi."],
        ["LightGBM", "Rekomendasi aksi dari enam fitur tabular bisnis.", "Ringan, cepat untuk inference lokal, cocok untuk fitur tabular, dan mudah dievaluasi/retrain batch (Ke et al., 2017)."],
        ["LLM agent", "Tidak digunakan sebagai pengambil keputusan harga.", "Berguna untuk bahasa dan ekstraksi, tetapi keluaran generatif tidak menjadi otoritas business rule."],
        ["Gemini structured output", "NLU terstruktur dan respons natural.", "Schema membantu ekstraksi lebih terprediksi; keluaran tetap divalidasi aplikasi (Google AI, 2026)."]
    ], [1.45, 2.05, 3.0])

    h("5. Metodologi Data dan Pelatihan Model")
    h("5.1 Dataset bootstrap", 2)
    para("Pada tahap MVP, LightGBM dilatih menggunakan 5.000 skenario negosiasi sintetis berbasis heuristik bisnis. Data sintetis digunakan secara transparan sebagai bootstrap karena log negosiasi riil UMKM belum cukup untuk pelatihan yang representatif. Dataset bukan klaim perilaku pelanggan riil dan tidak digunakan untuk menyimpulkan dampak bisnis populasi.")
    add_table(doc, ["Komponen", "Spesifikasi yang direproduksi"], [
        ["Jumlah data", "5.000 skenario; 4.000 train dan 1.000 test."],
        ["Pembagian", "Stratified train-test split 80:20; random_state=42."],
        ["Noise", "Controlled label noise 8% untuk menghindari dataset heuristik yang terlalu deterministik."],
        ["Fitur", "margin_pct, stock_ratio, customer_loyalty, discount_requested_pct, hour_of_day, is_peak_hour."],
        ["Label", "bonus, counter_offer, discount, hold_price."],
        ["Artefak", "ml/generate_synthetic_data.py, scoring_model.pkl, training_metadata.json."]
    ], [1.55, 4.95])
    h("5.2 Proses pembuatan dataset sintetis", 2)
    para("Dataset dibuat secara terprogram menggunakan skrip Python ml/generate_synthetic_data.py dengan numpy dan pandas, bukan melalui pengisian manual ataupun scraping data pelanggan. Generator memakai numpy.random.default_rng(42), sehingga eksperimen dapat diulang dan menghasilkan urutan data yang konsisten pada versi kode yang sama. Setelah fitur dibangkitkan, fungsi label_action memberi satu label aksi untuk setiap baris, lalu hasilnya ditulis menjadi ml/data/synthetic_negotiation_data.csv.")
    para("Dasar heuristik pada MVP disusun tim sebagai aturan bisnis awal yang masuk akal untuk konteks negosiasi, misalnya ruang margin lebih besar memberi kelonggaran lebih besar dan stok menipis mengurangi insentif memberi diskon. Heuristik ini bukan hasil wawancara/panel pakar maupun klaim representasi statistik UMKM Indonesia. Karena itu, ia disebut bootstrap dan divalidasi ulang melalui pilot sebelum dipakai untuk rekomendasi kebijakan yang lebih luas.")
    para("Data percakapan publik tidak digunakan untuk melatih tindakan harga pada MVP, karena label dan konteks margin/stoknya tidak identik dengan kondisi toko mitra. Jika digunakan pada fase berikutnya, data publik hanya boleh dipakai untuk eksperimen pemahaman bahasa setelah peninjauan lisensi, relevansi bahasa Indonesia, dan risiko privasi. Sumber pembelajaran harga yang dituju setelah MVP adalah log negosiasi pilot yang berizin dan dianonimkan.")
    h("5.3 Pembangkitan enam fitur", 2)
    add_table(doc, ["Fitur", "Cara dibangkitkan", "Makna operasional"], [
        ["margin_pct", "Uniform 0,05-0,50", "(price - floor_price) / price; ruang diskon relatif."],
        ["stock_ratio", "Uniform 0,00-1,00", "Stok ternormalisasi; 1 berarti stok relatif penuh."],
        ["customer_loyalty", "Beta(2,5)", "Membuat mayoritas skenario pelanggan baru, dengan sebagian kecil pelanggan lebih loyal."],
        ["discount_requested_pct", "Uniform 0,00-0,60", "Persentase diskon yang diminta dari harga awal."],
        ["hour_of_day", "Integer 0-23", "Jam percakapan saat skenario terjadi."],
        ["is_peak_hour", "Turunan hour_of_day", "1 untuk pukul 18-21; 0 untuk jam lain."]
    ], [1.5, 2.1, 2.9])
    h("5.4 Aturan penentuan label aksi", 2)
    para("Untuk setiap baris, generator menghitung ratio = discount_requested_pct / max(margin_pct, 0,05), stock_scarcity = 1 - stock_ratio, dan willingness = 0,35 x loyalty + 0,25 x stock_ratio - 0,10 x is_peak_hour + 0,20 x (1 - min(ratio, 1)). Aturan ini membuat hubungan antarfitur dapat dijelaskan dan diaudit sebelum model belajar dari pola tersebut.")
    add_table(doc, ["Kondisi heuristik", "Label yang dibangkitkan", "Rasional bisnis awal"], [
        ["ratio > 1,8", "counter_offer", "Permintaan diskon terlalu besar dibanding ruang margin; beri alternatif yang lebih aman."],
        ["1,0 < ratio <= 1,8", "hold_price", "Permintaan telah melampaui ruang margin relatif."],
        ["ratio <= 1,0; willingness > 0,42; stock_scarcity < 0,5; loyalty > 0,35", "bonus", "Pelanggan loyal dan stok cukup dapat diberi nilai tambah, bukan selalu diskon tunai."],
        ["ratio <= 1,0; willingness > 0,28", "discount", "Ada kelonggaran bisnis menurut sinyal gabungan."],
        ["Lainnya", "hold_price", "Menjaga harga dasar ketika kelonggaran tidak cukup."]
    ], [2.45, 1.25, 2.8])
    h("5.5 Controlled label noise 8%", 2)
    para("Sesudah label dasar ditetapkan, setiap baris memiliki peluang 0,08 untuk mengganti labelnya secara acak menjadi salah satu dari empat kelas: hold_price, discount, bonus, atau counter_offer. Teknik ini adalah random label replacement sederhana, bukan augmentasi bahasa atau perubahan fitur. Tujuannya agar dataset berbasis aturan tidak menjadi pemetaan yang terlalu sempurna; namun noise ini tidak menjadikan data sintetis setara dengan data lapangan. Semua metrik tetap dilaporkan sebagai metrik hold-out sintetis.")
    h("5.6 Prosedur pelatihan dan evaluasi", 2)
    para("Generator data menggunakan random number generator dengan seed 42. Label dibentuk oleh kombinasi margin, kelangkaan stok, loyalitas, diskon yang diminta, dan jam ramai; kemudian LightGBMClassifier dilatih secara offline. Saat live, model hanya melakukan inference. Retraining masa depan dilakukan secara batch dari negotiation_logs yang telah dianonimkan dan disetujui pemilik usaha, setelah kualitas dan volume data mencukupi.")
    add_table(doc, ["Metrik validasi internal", "Nilai"], [
        ["Accuracy", "0,892 (89,2%)"],
        ["Macro precision", "0,8912"],
        ["Macro recall", "0,8051"],
        ["Macro F1-score", "0,8320"],
        ["Waktu pelatihan", "2,6219 detik pada artefak eksperimen"],
    ], [3.4, 3.1])
    para("Batas interpretasi: angka di atas adalah evaluasi internal pada hold-out data sintetis, bukan akurasi terhadap transaksi UMKM nyata. Evaluasi pilot berikutnya harus melaporkan Word Error Rate voice note, akurasi intent berlabel manual, latency respons, serta conversion rate dari skenario uji yang terdokumentasi.")
    h("5.7 Interpretasi hasil dan reproduksibilitas", 2)
    para("Hasil eksperimen menunjukkan model dapat mempelajari pola dari generator yang sama pada test split terpisah. Ini bukan bukti bahwa model telah memahami semua perilaku tawar-menawar masyarakat Indonesia. Kelas bonus memiliki recall hold-out 0,4697, lebih rendah daripada kelas lain; hal ini konsisten dengan dukungan kelas bonus yang lebih kecil pada test set (66 sampel). Oleh sebab itu, confidence model dipakai sebagai sinyal dan keputusan tetap dibatasi hard guardrail. Artefak scoring_model.pkl, label_encoder.pkl, dan training_metadata.json disimpan untuk evaluasi ulang melalui ml/evaluate_model.py.")
    para("Proposal tidak menyajikan baseline rule-based sebagai angka pembanding karena label dataset sendiri dibentuk oleh heuristik yang sama; membandingkan keduanya akan menghasilkan evaluasi yang bias dan tidak bermakna. Baseline yang valid untuk pilot adalah kebijakan operasional sebelum LARISKA, misalnya median waktu respons, tingkat percakapan terselesaikan, dan tingkat checkout dalam periode pembanding yang didefinisikan sejak awal.")
    h("5.8 Model card ringkas", 2)
    add_table(doc, ["Elemen", "Keterangan"], [
        ["Tujuan", "Merekomendasikan satu dari empat aksi negosiasi; bukan menetapkan harga akhir."],
        ["Input", "Enam fitur tabular: margin, stok relatif, loyalitas, diskon diminta, jam, dan jam ramai."],
        ["Output", "bonus, counter_offer, discount, atau hold_price beserta confidence bila tersedia."],
        ["Batasan", "Terlatih dari data sintetis; tidak boleh dipakai sebagai bukti perilaku pasar atau kebijakan diskon tanpa pilot."],
        ["Kontrol", "Guardrail floor price dan stok merupakan otoritas final; human handover menangani kasus sensitif."],
        ["Monitoring masa depan", "Drift fitur/label, distribusi keputusan, penerimaan negosiasi, insiden guardrail, dan audit fairness."]
    ], [1.55, 4.95])

    h("6. Guardrail, Keamanan, dan Keandalan")
    bullets([
        "Floor price dan stok diperiksa pada kode deterministik, bukan hanya melalui instruksi prompt.",
        "Harga dari client tidak dipercaya pada jalur order; backend menghitung harga dari katalog dan menolak diskon yang melanggar batas aman.",
        "Webhook Meta dan Midtrans dicatat secara idempoten untuk mengurangi risiko event ganda; payment status bukan input operator dashboard.",
        "Status operasional order memiliki alur pending, paid, shipped, completed, atau cancelled. Dashboard tidak dapat memaksa pending menjadi paid.",
        "Human handover dipicu untuk permintaan admin eksplisit atau komplain serius, bukan untuk negosiasi normal.",
        "Data percakapan dan transaksi harus dikelola dengan consent, retensi minimum, akses berbasis peran, dan prinsip UU Pelindungan Data Pribadi (Republik Indonesia, 2022)."
    ])
    h("6.1 Tabel risiko dan mitigasi", 2)
    add_table(doc, ["Risiko", "Mitigasi MVP", "Penguatan berikutnya"], [
        ["LLM mengarang fakta", "Fakta katalog dan harga dijawab dari Supabase; LLM dibatasi pada bahasa.", "Evaluasi grounding dan dataset uji regresi."],
        ["Tawaran merugikan", "Floor price diperiksa server-side dan tersimpan pada negotiation log.", "Konfigurasi kebijakan per UMKM dan dashboard approval."],
        ["Webhook ganda", "Event Meta/Midtrans dideduplikasi dan diproses idempoten.", "Dead-letter queue dan observability produksi."],
        ["Audio tidak jelas", "Whisper small dengan fallback base/Gemini Audio; minta klarifikasi bila perlu.", "Benchmark WER berlabel dan perbaikan vocabulary domain."],
        ["Data pribadi", "Prinsip consent, retensi minimum, role-based access sebagai kebijakan target.", "DPIA, enkripsi terkelola, dan SOP penghapusan/akses data."]
    ], [1.4, 2.8, 2.3])

    h("7. Implementasi MVP dan Bukti Kesiapan")
    add_table(doc, ["Kemampuan", "Implementasi saat ini"], [
        ["Katalog", "CRUD produk, SKU, satuan, spesifikasi yang mudah diisi, floor price, stok, reorder point, dan alias pencarian."],
        ["Chat commerce", "Greeting dengan interactive list kategori, tanya katalog deterministik, rekomendasi, negosiasi, checkout, dan CTA pembayaran."],
        ["Voice intelligence", "Penerimaan voice note WhatsApp, transkripsi Whisper, normalisasi melalui NLU structured output."],
        ["Pembayaran", "Midtrans Sandbox QRIS/payment link, payment webhook, pembayaran success, order paid, reservasi dan pengurangan stok."],
        ["Operasional", "Dashboard produk, pelanggan, order, payment snapshot, dan transisi pengiriman yang aman."],
        ["Keandalan", "Fallback respons saat LLM gagal, retry client Gemini, deduplikasi webhook, dan test suite backend offline."],
    ], [1.55, 4.95])
    h("7.1 Fondasi data dan keterlacakan", 2)
    para("Supabase PostgreSQL menyimpan customers, products, conversations, messages, orders, payments, negotiation_logs, customer_memory, recommendations, business_insights, dan inventory_logs. Fondasi transaksi ditambah dengan order_items, product_variants, inventory_reservations, payment_events, serta webhook_events. Struktur ini memisahkan katalog, percakapan, keputusan negosiasi, pembayaran, dan pergerakan stok sehingga dashboard dapat menjelaskan bukan hanya hasil akhir, tetapi proses yang menghasilkannya.")
    para("Bukti integrasi yang telah diuji pada sandbox: webhook Midtrans settlement mengubah payment menjadi success dan order menjadi paid, mengonfirmasi reservasi, lalu mencatat pengurangan stok dari 19 menjadi 18 untuk kuantitas satu. Status pemenuhan berikutnya dilakukan operator melalui paid -> shipped -> completed. Bukti ini adalah hasil skenario sandbox, bukan volume transaksi produksi.")
    h("7.2 Strategi pengujian", 2)
    para("Suite backend offline saat ini memverifikasi kontrak Pydantic, guardrail harga, checkout/payment, handover, klien WhatsApp, katalog, dan pipeline tanpa menguras kuota LLM. Pengujian live dipisahkan karena bergantung pada kredensial dan kuota eksternal. Pendekatan ini memberi keseimbangan antara regresi deterministik yang dapat diulang dan smoke test integrasi yang hemat kuota.")
    h("7.3 Status bukti", 2)
    add_table(doc, ["Area", "Status yang dapat diklaim", "Tidak diklaim"], [
        ["WhatsApp", "Webhook dan alur teks/voice note diuji pada environment Meta yang sesuai konfigurasi aplikasi.", "Layanan publik tanpa batas penerima uji sebelum aplikasi/nomor diproduksikan."],
        ["Pembayaran", "Midtrans Sandbox membuktikan alur payment link, settlement webhook, paid, dan stok.", "Validasi transaksi finansial produksi atau pendapatan riil."],
        ["AI", "Pipeline dan regresi offline dapat diuji; metrik LightGBM tersedia pada hold-out sintetis.", "Akurasi model terhadap populasi UMKM nyata."],
        ["Dashboard", "Mengelola katalog, pelanggan, order, payment snapshot, dan status pemenuhan.", "Kemampuan SaaS multi-tenant produksi pada MVP saat ini."]
    ], [1.35, 3.0, 2.15])

    h("8. Model Bisnis dan Strategi Go-to-Market")
    para("LARISKA diposisikan sebagai B2B SaaS untuk UMKM dengan volume chat tinggi, terutama F&B, fashion, aksesori, kebutuhan harian, elektronik, dan sparepart. Kanal WhatsApp dipilih karena dekat dengan pola kerja pelaku UMKM dan pelanggan Indonesia.")
    add_table(doc, ["Tahap", "Strategi"], [
        ["Pilot", "Onboarding satu UMKM per kategori prioritas, mengukur waktu respons, pertanyaan terselesaikan, dan conversion funnel."],
        ["Monetisasi", "Langganan bulanan bertingkat berdasarkan volume percakapan/fitur; opsi biaya transaksi hanya setelah struktur biaya dan consent jelas."],
        ["Retensi", "Business Copilot, katalog terstruktur, riwayat pelanggan, dan log negosiasi membuat nilai bertambah dari penggunaan sehari-hari."],
        ["Ekspansi", "Multi-tenant workspace, nomor WhatsApp per UMKM, role admin, dan model retraining batch sebagai fase setelah MVP tervalidasi."],
    ], [1.3, 5.2])
    h("8.1 Batas model bisnis pada MVP", 2)
    para("MVP membuktikan alur satu operasional usaha dan belum mengklaim bahwa multi-tenant SaaS telah selesai diproduksikan. Perluasan ke banyak UMKM membutuhkan isolasi tenant, peran pengguna, pengelolaan nomor WhatsApp bisnis, billing, dan kebijakan data yang tervalidasi. Menyatakan batas ini secara jelas menjaga proposal tetap feasible dan dapat dipertanggungjawabkan.")

    h("9. Diferensiasi dan Dampak")
    para("Diferensiasi LARISKA bukan sekadar koneksi ke LLM. Nilainya berada pada orchestration: suara dan bahasa informal diterjemahkan menjadi data terstruktur; keputusan harga dipagari; fakta katalog dibaca dari database; dan transaksi kembali menjadi data untuk operasional. Hal ini membuat solusi relevan untuk UMKM yang kewalahan melayani banyak pelanggan sekaligus tanpa menyerahkan kontrol margin kepada model bahasa.")
    para("Moat yang realistis dibangun melalui konfigurasi guardrail yang spesifik per UMKM, log interaksi yang berizin, serta evaluasi dan retraining batch berbasis data lapangan. Pernyataan ini adalah roadmap, bukan klaim bahwa sistem telah melakukan online learning pada MVP.")
    h("9.1 Indikator dampak pilot", 2)
    add_table(doc, ["Dimensi", "Indikator", "Cara membaca"], [
        ["Layanan", "Median waktu respons dan tingkat percakapan terselesaikan", "Bandingkan baseline sebelum/sesudah dengan periode yang sama."],
        ["Kualitas AI", "Intent accuracy, WER voice note, hallucination/fallback rate", "Dinilai dari sampel berlabel manual dan log audit."],
        ["Bisnis", "Checkout initiation, payment success, negotiation acceptance", "Bukan klaim kausal tanpa desain pilot yang memadai."],
        ["Keamanan", "Pelanggaran floor price, mismatch stok, event webhook duplikat", "Targetnya nol pelanggaran guardrail dan investigasi setiap insiden."]
    ], [1.25, 2.7, 2.55])
    h("9.2 Kontribusi inovasi", 2)
    bullets([
        "Mengubah chat commerce dari jawaban terpisah menjadi alur keputusan yang menyambungkan discovery, negosiasi, checkout, dan fulfillment.",
        "Menggunakan AI secara proporsional: model generatif membantu bahasa, sedangkan keputusan berisiko dikendalikan oleh aturan dan data operasional.",
        "Menyediakan jejak audit untuk harga, stok, pembayaran, serta outcome negosiasi agar pemilik UMKM tetap memegang kontrol." 
    ])

    h("10. Rencana Evaluasi dan Demo")
    para("Demo dibuat sebagai bukti kerja ujung-ke-ujung, bukan sekadar tampilan dashboard. Urutan yang disarankan adalah:")
    for i, step in enumerate([
        "Kirim 'Halo' dan tunjukkan interactive list kategori dari katalog aktif.",
        "Pilih kategori, lalu tanya detail, stok, dan harga salah satu produk.",
        "Kirim voice note untuk pertanyaan produk dan tampilkan hasil transkripsi serta balasan.",
        "Ajukan negosiasi multi-unit; tampilkan counter-offer yang tidak melanggar floor price.",
        "Checkout dan selesaikan QRIS Sandbox; tunjukkan order paid, payment success, dan stok berkurang.",
        "Tandai order shipped lalu completed dari dashboard untuk menunjukkan loop operasional."
    ], start=1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    para("Sebelum presentasi, tim perlu merekam demo cadangan yang menampilkan timestamp, terminal/backend, dan status dashboard. Metrik yang dikumpulkan untuk pilot adalah latency respons, WER voice note, akurasi intent pada sampel berlabel manual, tingkat checkout dari skenario uji, serta insiden guardrail.")
    h("10.1 Naskah pesan kunci untuk demonstrasi", 2)
    bullets([
        "Saat produk ditanya, LARISKA membaca katalog nyata: ini mencegah jawaban harga, ukuran, dan stok yang tidak konsisten.",
        "Saat pelanggan menawar, AI tidak bebas menentukan harga: LightGBM merekomendasikan aksi dan guardrail memastikan lantai harga tidak dilanggar.",
        "Saat pembayaran berhasil, sistem tidak sekadar menampilkan status; webhook mengubah state order dan mencatat dampak ke inventori.",
        "Saat pelanggan lebih nyaman berbicara, voice note tetap masuk ke alur keputusan yang sama melalui transkripsi."
    ])

    h("11. Roadmap Setelah MVP")
    add_table(doc, ["Horizon", "Prioritas"], [
        ["0-3 bulan", "Pilot terkontrol, baseline evaluasi nyata, perbaikan katalog dan UX percakapan, SOP handover."],
        ["3-6 bulan", "Tenant/workspace per UMKM, role-based access, nomor WhatsApp per tenant, observability dan alerting."],
        ["6-12 bulan", "Retraining batch berizin dari log teranonymisasi, evaluasi fairness, dan integrasi logistik/kurir."],
    ], [1.35, 5.15])
    h("11.1 Tahap tata kelola data", 2)
    para("Sebelum retraining dari data pilot, tim menetapkan mekanisme persetujuan pemilik usaha, minimisasi data, penghilangan identitas langsung pada dataset analitik, pemisahan akses operasional dan analitik, evaluasi drift, serta review manual atas label/hasil yang berdampak pada kebijakan harga. Roadmap ini menjaga peningkatan model tidak mengorbankan kepercayaan pelaku UMKM dan pelanggan.")

    h("Kesimpulan")
    para("LARISKA AI menawarkan pendekatan praktis untuk digitalisasi penjualan UMKM: bukan mengganti pemilik usaha dengan LLM, melainkan memberi sistem yang mampu mengelola percakapan, menjaga aturan bisnis, menutup transaksi, dan menyediakan jejak operasional. Arsitektur hibrida membuat solusi lebih aman untuk negosiasi harga, lebih dapat diuji untuk kompetisi, dan lebih siap dikembangkan menjadi produk nyata.")

    h("Daftar Pustaka")
    refs = [
        "Badan Pusat Statistik. (2025). Profil Industri Mikro dan Kecil 2024. https://www.bps.go.id/assets/publication/2025/09/16/a83f105e49377d0a7434e62a/profil-industri-mikro-dan-kecil-2024.html",
        "Badan Pusat Statistik. (2024). Statistik Indonesia 2024. https://www.bps.go.id/id/publication/2024/02/28/c1bacde03256343b2bf769b0/statistik-indonesia-2024.html",
        "Republik Indonesia. (2022). Undang-Undang Nomor 27 Tahun 2022 tentang Pelindungan Data Pribadi. https://peraturan.bpk.go.id/Details/229798/uu-no-27-tahun-2022",
        "Ke, G., et al. (2017). LightGBM: A Highly Efficient Gradient Boosting Decision Tree. Advances in Neural Information Processing Systems 30.",
        "Radford, A., et al. (2023). Robust Speech Recognition via Large-Scale Weak Supervision. Proceedings of ICML 2023.",
        "Meta. WhatsApp Cloud API Documentation. https://developers.facebook.com/docs/whatsapp/cloud-api/ (diakses Agustus 2026).",
        "Google AI. (2026). Structured outputs - Gemini API. https://ai.google.dev/gemini-api/docs/structured-output (diakses Agustus 2026).",
        "Artefak internal LARISKA AI: ml/generate_synthetic_data.py dan ml/model_artifacts/training_metadata.json (eksperimen ter-reproduksi, 2026)."
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "LARISKA AI - Proposal Inovasi Final"
    doc.core_properties.author = "Kelompok iki"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
